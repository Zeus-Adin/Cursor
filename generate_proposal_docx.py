#!/usr/bin/env python3
"""Generate StacksPot Growth Acceleration Proposal as a Word document."""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(0x1A, 0x36, 0x5D)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
DARK = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x4A, 0x55, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_FILL = "EEF2F7"
HEADER_FILL = "1A365D"
ALT_ROW = "F7FAFC"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)
    return run


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, color="CBD5E1", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_cell_text(cell, text, bold=False, size=9, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, before=2, after=2, line=1.0)
    add_text(p, text, size=size, bold=bold, color=color)


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=16 if level == 1 else 10, after=6, line=1.15)
    if level == 1:
        add_text(p, text, name="Calibri", size=15, bold=True, color=NAVY)
    else:
        add_text(p, text, name="Calibri", size=12, bold=True, color=ACCENT)
    return p


def add_body(doc, text, after=8):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=after, line=1.15)
    add_text(p, text, size=11, color=DARK)
    return p


def add_rich_body(doc, parts, after=8):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=after, line=1.15)
    for text, kwargs in parts:
        add_text(p, text, **kwargs)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=3, line=1.1)
    if p.runs:
        p.runs[0].text = ""
    add_text(p, text, size=11, color=DARK)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, before=0, after=3, line=1.1)
    if p.runs:
        p.runs[0].text = ""
    add_text(p, text, size=11, color=DARK)
    return p


def is_total_row(row):
    joined = " ".join(str(x) for x in row).lower()
    return any(k in joined for k in ("total", "5-month", "program", "23 total", "lifecycle total"))


def make_table(doc, headers, rows, col_widths=None, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=font_size, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, HEADER_FILL)
        set_cell_borders(cell, color="1A365D", size="4")

    for r_idx, row in enumerate(rows):
        total = is_total_row(row)
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            if total:
                set_cell_text(cell, str(val), bold=True, size=font_size, color=NAVY, align=align)
                shade_cell(cell, LIGHT_FILL)
            else:
                set_cell_text(cell, str(val), bold=False, size=font_size, color=DARK, align=align)
                if r_idx % 2 == 1:
                    shade_cell(cell, ALT_ROW)
            set_cell_borders(cell)

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=6, line=1.0)
    return table


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=4, line=1.0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2B6CB0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        set_cell_text(c0, label, bold=True, size=10, color=WHITE)
        shade_cell(c0, HEADER_FILL)
        set_cell_borders(c0, color="1A365D")
        set_cell_text(c1, value, size=10, color=DARK)
        set_cell_borders(c1)
        shade_cell(c1, ALT_ROW if i % 2 else "FFFFFF")
        c0.width = Inches(2.4)
        c1.width = Inches(4.1)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=12, after=4, line=1.1)
    add_text(title, "StacksPot Growth Acceleration Proposal", name="Calibri", size=22, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, before=0, after=10, line=1.1)
    add_text(
        subtitle,
        "A Sponsored Yield Boosting Initiative for Ecosystem Growth",
        name="Calibri",
        size=12,
        italic=True,
        color=ACCENT,
    )
    add_horizontal_line(doc)

    add_meta_table(
        doc,
        [
            ("Submitted By", "StacksPot Team"),
            ("Submitted To", "Stacks Endowment Team"),
            ("Program Duration", "5 months · 10 cycles"),
            ("Cycle Length", "2 weeks"),
            ("Pots per Cycle", "23"),
            ("Pots per Month", "46 (2 cycles)"),
            ("Total Pots (Program)", "230"),
            ("Requested Allocation", "500,000 STX"),
            ("Capital Model", "Reusable — returned to sponsor after each cycle"),
            ("Date", "July 2026"),
        ],
    )
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=4, after=4)

    add_heading_styled(doc, "1. Objective")
    add_body(
        doc,
        "Accelerate Stacks user adoption through sponsored pot campaigns that boost rewards, deepen STX activity, and introduce more participants to Bitcoin-secured DeFi on Stacks.",
    )

    add_heading_styled(doc, "2. Executive Summary")
    add_body(
        doc,
        "StacksPot is a decentralized, yield-powered community pot protocol on Stacks. Users join transparent reward pools while preserving principal and earning Bitcoin-secured DeFi yield.",
    )
    add_rich_body(
        doc,
        [
            ("Early growth is constrained less by product readiness than by ", {"size": 11}),
            ("reward perception", {"size": 11, "bold": True}),
            ("—users weigh expected upside against commitment size and wait time. To close that gap, StacksPot requests a ", {"size": 11}),
            ("500,000 STX", {"size": 11, "bold": True}),
            (" Endowment sponsorship allocation for a ", {"size": 11}),
            ("5-month Growth Accelerator", {"size": 11, "bold": True}),
            (".", {"size": 11}),
        ],
    )
    add_rich_body(
        doc,
        [
            ("Sponsorship capital is ", {"size": 11}),
            ("delegated into pots each 2-week cycle", {"size": 11, "bold": True}),
            (", then ", {"size": 11}),
            ("returned to the sponsor address", {"size": 11, "bold": True}),
            (" when the cycle closes. The ", {"size": 11}),
            ("same 500,000 STX", {"size": 11, "bold": True}),
            (" is redeployed across all ", {"size": 11}),
            ("10 cycles", {"size": 11, "bold": True}),
            (". The Endowment’s capital ask does not multiply with cycle count.", {"size": 11}),
        ],
    )
    make_table(
        doc,
        ["Scope", "Cycles", "Pots", "Capital Required"],
        [
            ("Per cycle", "1", "23", "500,000 STX (delegated)"),
            ("Per month", "2", "46", "500,000 STX (same capital, reused)"),
            ("5-month program", "10", "230", "500,000 STX (same capital, reused)"),
        ],
        col_widths=[1.8, 1.0, 1.0, 2.8],
    )

    add_heading_styled(doc, "3. How Sponsorship Capital Works")
    add_numbered(doc, "Endowment allocates 500,000 STX to the sponsor address.")
    add_numbered(doc, "That STX is delegated across the 23 pots in the active cycle.")
    add_numbered(doc, "When the cycle ends, delegated amounts are returned to the sponsor address.")
    add_numbered(doc, "The same STX is redeployed for the next cycle.")
    add_numbered(doc, "This repeats for 10 cycles over 5 months.")
    add_body(doc, "One allocation powers the full campaign without requiring new capital each cycle.")

    add_heading_styled(doc, "4. Boosting Campaign Structure")
    add_body(
        doc,
        "This table is the full breakdown of the sponsored campaign. All figures use the reusable 500,000 STX allocation across 10 cycles (5 months).",
    )
    add_rich_body(
        doc,
        [
            ("(*) Pot type: ", {"size": 10, "italic": True, "color": GRAY}),
            ("JackPot or Sequential, as noted in Qty.", {"size": 10, "italic": True, "color": GRAY}),
        ],
        after=6,
    )

    make_table(
        doc,
        [
            "Qty of Pots / Cycle (*)",
            "Entry Min.",
            "Min. Participants",
            "Pot Target",
            "Total Tx",
            "Total Stacked",
            "Total Participants",
            "Awarded Boost",
        ],
        [
            ("10 (JackPot)", "25 STX", "25", "625 STX", "2,900", "62,500 STX", "2,500", "10,000 STX / pot"),
            ("12 (JackPot)", "50 STX", "50", "2,500 STX", "6,480", "300,000 STX", "6,000", "25,000 STX / pot"),
            ("1 (Sequential)", "350 STX", "10", "3,500 STX", "140", "35,000 STX", "100", "100,000 STX / pot"),
            ("23 total", "—", "—", "—", "9,520", "397,500 STX", "8,600", "500,000 STX / cycle"),
        ],
        col_widths=[1.15, 0.7, 0.95, 0.75, 0.6, 0.9, 0.9, 1.1],
        font_size=8,
    )

    add_heading_styled(doc, "Tier Notes", level=2)
    make_table(
        doc,
        ["Tier", "Pot Type", "Pots / Cycle", "Entry Min.", "Min. Part. / Pot", "Pot Target", "Boost / Pot", "Tier Capital"],
        [
            ("A", "JackPot", "10", "25 STX", "25", "625 STX", "10,000 STX", "100,000 STX"),
            ("B", "JackPot", "12", "50 STX", "50", "2,500 STX", "25,000 STX", "300,000 STX"),
            ("C", "Sequential", "1", "350 STX", "10", "3,500 STX", "100,000 STX", "100,000 STX"),
            ("Total", "—", "23", "—", "—", "—", "—", "500,000 STX"),
        ],
        col_widths=[0.55, 0.9, 0.85, 0.75, 0.95, 0.85, 0.95, 1.0],
        font_size=8,
    )

    note = doc.add_paragraph()
    set_paragraph_spacing(note, before=0, after=4, line=1.1)
    add_text(note, "Tier B: ", size=9, italic=True, bold=True, color=GRAY)
    add_text(
        note,
        "12 JackPot pots share 300,000 STX boost capital per cycle → 25,000 STX / pot. ",
        size=9,
        italic=True,
        color=GRAY,
    )
    add_text(note, "Pot Target: ", size=9, italic=True, bold=True, color=GRAY)
    add_text(note, "Entry Min. × Min. Participants. ", size=9, italic=True, color=GRAY)
    add_text(note, "Total Stacked: ", size=9, italic=True, bold=True, color=GRAY)
    add_text(note, "Pot Target × Qty × 10 cycles (participant STX at min fill). ", size=9, italic=True, color=GRAY)
    add_text(note, "Total Participants: ", size=9, italic=True, bold=True, color=GRAY)
    add_text(note, "Min. Participants × Qty × 10 cycles. ", size=9, italic=True, color=GRAY)
    add_text(note, "Total Tx: ", size=9, italic=True, bold=True, color=GRAY)
    add_text(
        note,
        "lifecycle (4 × pots × 10) + joins. Awarded Boost / Tier Capital = reusable sponsor STX (not in Total Stacked).",
        size=9,
        italic=True,
        color=GRAY,
    )

    add_heading_styled(doc, "How the Totals Are Calculated", level=2)
    make_table(
        doc,
        ["Tier", "Pot Target", "Total Stacked (10 cycles)", "Lifecycle Txs", "Join Txs", "Total Tx", "Total Participants"],
        [
            ("A — 10 × 25 STX × 25 min", "625", "62,500 STX", "400", "2,500", "2,900", "2,500"),
            ("B — 12 × 50 STX × 50 min", "2,500", "300,000 STX", "480", "6,000", "6,480", "6,000"),
            ("C — 1 × 350 STX × 10 min", "3,500", "35,000 STX", "40", "100", "140", "100"),
            ("Program", "—", "397,500 STX", "920", "8,600", "9,520", "8,600"),
        ],
        col_widths=[1.8, 0.8, 1.4, 1.0, 0.85, 0.8, 1.2],
        font_size=8,
    )

    add_heading_styled(doc, "Transaction Accounting", level=2)
    make_table(
        doc,
        ["Transaction", "Per Pot / Per Join", "Tier A", "Tier B", "Tier C", "Program"],
        [
            ("Pot deploy", "1 / pot", "100", "120", "10", "230"),
            ("Pot activate", "1 / pot", "100", "120", "10", "230"),
            ("Pot start", "1 / pot", "100", "120", "10", "230"),
            ("Pot close", "1 / pot", "100", "120", "10", "230"),
            ("Participant join", "1 / participant", "2,500", "6,000", "100", "8,600"),
            ("Total txs", "—", "2,900", "6,480", "140", "9,520"),
        ],
        col_widths=[1.4, 1.3, 1.0, 1.0, 0.9, 1.0],
    )

    add_heading_styled(doc, "5. Expected Outcomes")
    make_table(
        doc,
        ["Outcome", "Target"],
        [
            ("Sponsored pots delivered", "230 (23 / cycle × 10 cycles)"),
            ("Minimum participants (program)", "8,600 (2,500 + 6,000 + 100)"),
            ("Participant STX stacked (minimum)", "397,500 STX (62,500 + 300,000 + 35,000)"),
            ("On-chain transactions (minimum)", "9,520 (920 lifecycle + 8,600 joins)"),
            ("Monthly throughput", "46 pots / month · 1,720 min. participants · 2 cycles"),
            ("Capital efficiency", "500,000 STX committed once; returned and reused every cycle"),
            ("Ecosystem impact", "Higher STX deposits, wallet activity, and Bitcoin DeFi awareness on Stacks"),
        ],
        col_widths=[2.4, 4.1],
    )

    add_heading_styled(doc, "6. Reporting")
    add_rich_body(
        doc,
        [
            ("Monthly: ", {"size": 11, "bold": True}),
            ("participants, STX deposited, pots created (target 46/month), capital deployed and returned each cycle, and transaction counts (deploy / activate / start / close / joins).", {"size": 11}),
        ],
    )
    add_rich_body(
        doc,
        [
            ("Final (Month 5): ", {"size": 11, "bold": True}),
            ("full 10-cycle performance, retention, cumulative STX activity, capital custody summary, and recommendations.", {"size": 11}),
        ],
    )

    add_heading_styled(doc, "7. Conclusion")
    add_rich_body(
        doc,
        [
            ("StacksPot requests a ", {"size": 11}),
            ("500,000 STX", {"size": 11, "bold": True}),
            (" reusable sponsorship allocation to run ", {"size": 11}),
            ("23 boosted pots every 2-week cycle", {"size": 11, "bold": True}),
            (" for ", {"size": 11}),
            ("5 months", {"size": 11, "bold": True}),
            (" (", {"size": 11}),
            ("230 pots", {"size": 11, "bold": True}),
            (" program-wide). Because delegated capital returns to the sponsor after each cycle, one allocation funds the entire campaign—attracting users, supporting communities, and expanding Bitcoin-secured DeFi activity on Stacks.", {"size": 11}),
        ],
    )

    add_horizontal_line(doc)
    add_meta_table(
        doc,
        [
            ("Prepared by", "StacksPot Team"),
            ("Date", "July 2026"),
            ("Requested allocation", "500,000 STX (reusable — returned after each cycle)"),
            ("Campaign mix", "10 JackPot × 10k · 12 JackPot × 25k · 1 Sequential × 100k"),
        ],
    )

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(
        fp,
        "StacksPot Growth Acceleration Proposal  |  Confidential — For Stacks Endowment Review",
        size=8,
        color=GRAY,
        italic=True,
    )

    out = "/workspace/StacksPot-Growth-Acceleration-Proposal.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    build()
